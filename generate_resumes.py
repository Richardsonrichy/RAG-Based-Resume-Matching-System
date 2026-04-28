"""
generate_resumes.py
Generates 30+ diverse synthetic resumes as .docx files in data/resumes/
Run this ONCE before resume_rag.py to build your dataset.
"""

import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "data/resumes"

# ─── Resume templates ──────────────────────────────────────────────────────────

RESUMES = [
    {
        "name": "Aarav Mehta",
        "email": "aarav.mehta@gmail.com", "phone": "+91 9876501001",
        "summary": "Python backend developer with 6+ years of experience in building scalable REST APIs, microservices, and data pipelines.",
        "skills": "Python, Django, FastAPI, PostgreSQL, Redis, Docker, Kubernetes, AWS, Git, REST APIs",
        "experience": [
            ("Senior Backend Engineer", "DataSoft Pvt Ltd", "2021 - Present",
             ["Designed and deployed microservices using Django and FastAPI",
              "Reduced API response time by 40% via Redis caching",
              "Led migration of monolith to Kubernetes-based architecture"]),
            ("Backend Developer", "TechNova", "2018 - 2021",
             ["Built REST APIs serving 2M+ daily requests",
              "Integrated third-party payment gateways",
              "Implemented CI/CD pipelines using Jenkins"]),
        ],
        "education": "B.Tech in Computer Science, IIT Bombay (2014 - 2018)",
        "exp_years": 6,
    },
    {
        "name": "Priya Sharma",
        "email": "priya.sharma@outlook.com", "phone": "+91 9876502002",
        "summary": "Data Scientist with 4+ years specializing in NLP, machine learning, and recommendation systems.",
        "skills": "Python, Scikit-learn, TensorFlow, NLP, Pandas, NumPy, SQL, Spark, MLflow, Airflow",
        "experience": [
            ("Data Scientist", "Analytics Corp", "2022 - Present",
             ["Built NLP models for sentiment analysis with 92% accuracy",
              "Developed recommendation engine improving CTR by 25%",
              "Automated ML pipelines using Apache Airflow"]),
            ("Junior Data Analyst", "InfoSys Ltd", "2020 - 2022",
             ["Performed EDA on large datasets using Pandas",
              "Created dashboards using Power BI",
              "Wrote complex SQL queries for reporting"]),
        ],
        "education": "M.Sc in Data Science, IIT Delhi (2018 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Rohan Verma",
        "email": "rohan.verma@gmail.com", "phone": "+91 9876503003",
        "summary": "DevOps Engineer with 5+ years of experience in Azure, Kubernetes, Terraform, and CI/CD automation.",
        "skills": "Azure, Kubernetes, Docker, Terraform, Ansible, Jenkins, Python, Bash, Prometheus, Grafana, Git",
        "experience": [
            ("Senior DevOps Engineer", "CloudBase Systems", "2021 - Present",
             ["Managed Azure infrastructure for 100+ microservices",
              "Automated deployments using Terraform and Ansible",
              "Implemented monitoring with Prometheus and Grafana"]),
            ("DevOps Engineer", "SysTech", "2019 - 2021",
             ["Setup CI/CD pipelines reducing release cycle by 60%",
              "Dockerized legacy applications",
              "Managed on-prem to cloud migrations"]),
        ],
        "education": "B.Tech in Information Technology, NIT Trichy (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Sneha Iyer",
        "email": "sneha.iyer@gmail.com", "phone": "+91 9876504004",
        "summary": "Full Stack Developer with 3+ years of experience in React, Node.js, and MongoDB.",
        "skills": "React, Node.js, Express, MongoDB, TypeScript, HTML, CSS, REST APIs, Git, Docker",
        "experience": [
            ("Full Stack Developer", "WebBridge Pvt Ltd", "2022 - Present",
             ["Built responsive React dashboards for B2B clients",
              "Developed Node.js backend with Express and MongoDB",
              "Integrated Stripe payment APIs"]),
            ("Frontend Developer", "PixelWorks", "2021 - 2022",
             ["Developed reusable React component libraries",
              "Improved page load speed by 35% through lazy loading",
              "Collaborated on UI/UX design with Figma"]),
        ],
        "education": "B.Tech in Computer Science, VIT Vellore (2017 - 2021)",
        "exp_years": 3,
    },
    {
        "name": "Karthik Subramaniam",
        "email": "karthik.sub@gmail.com", "phone": "+91 9876505005",
        "summary": "Java backend developer with 7+ years of experience in Spring Boot, Kafka, and distributed systems.",
        "skills": "Java, Spring Boot, Kafka, PostgreSQL, Microservices, REST APIs, Docker, AWS, Maven, Git",
        "experience": [
            ("Lead Software Engineer", "FinTechPro", "2020 - Present",
             ["Architected distributed payment processing system using Kafka",
              "Reduced system downtime by 50% through circuit breakers",
              "Mentored team of 5 junior engineers"]),
            ("Software Engineer", "Wipro Technologies", "2017 - 2020",
             ["Developed Spring Boot microservices for banking domain",
              "Optimized SQL queries reducing report generation by 60%",
              "Integrated third-party SWIFT APIs"]),
        ],
        "education": "B.E in Computer Science, Anna University (2013 - 2017)",
        "exp_years": 7,
    },
    {
        "name": "Ananya Krishnan",
        "email": "ananya.k@gmail.com", "phone": "+91 9876506006",
        "summary": "Cloud Solutions Architect with 8+ years in AWS, GCP, and enterprise cloud migrations.",
        "skills": "AWS, GCP, Azure, Terraform, CloudFormation, Python, Kubernetes, Docker, VPC, IAM, Lambda",
        "experience": [
            ("Cloud Architect", "GlobalTech Solutions", "2019 - Present",
             ["Designed multi-cloud architecture for Fortune 500 clients",
              "Led AWS to GCP migration saving $2M annually",
              "Implemented Zero Trust security model across environments"]),
            ("Cloud Engineer", "CyberNet", "2016 - 2019",
             ["Deployed serverless applications using AWS Lambda",
              "Automated infrastructure provisioning with CloudFormation",
              "Managed IAM policies for 500+ users"]),
        ],
        "education": "M.Tech in Computer Science, IISc Bangalore (2012 - 2016)",
        "exp_years": 8,
    },
    {
        "name": "Vikram Nair",
        "email": "vikram.nair@gmail.com", "phone": "+91 9876507007",
        "summary": "Machine Learning Engineer with 5+ years building production ML models and LLM-based applications.",
        "skills": "Python, PyTorch, TensorFlow, LLM, RAG, LangChain, Kubernetes, MLflow, Kubeflow, SQL, Git",
        "experience": [
            ("ML Engineer", "AI Ventures", "2021 - Present",
             ["Fine-tuned LLMs for enterprise document Q&A systems",
              "Built RAG pipelines using LangChain and ChromaDB",
              "Deployed ML models to production using Kubeflow"]),
            ("Data Scientist", "InsightLabs", "2019 - 2021",
             ["Developed churn prediction model with 88% recall",
              "Built real-time fraud detection using streaming ML",
              "Automated feature engineering pipelines"]),
        ],
        "education": "M.Tech in AI, IIT Madras (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Meera Pillai",
        "email": "meera.pillai@gmail.com", "phone": "+91 9876508008",
        "summary": "Data Engineer with 4+ years building data lakes, ETL pipelines, and real-time streaming solutions.",
        "skills": "Python, Apache Spark, Kafka, Airflow, AWS Glue, Redshift, Snowflake, SQL, dbt, Terraform",
        "experience": [
            ("Senior Data Engineer", "DataFlow Inc", "2022 - Present",
             ["Built petabyte-scale data lake on AWS S3 and Glue",
              "Developed real-time streaming pipelines with Kafka and Spark",
              "Reduced ETL job failures by 70% via robust error handling"]),
            ("Data Engineer", "Analytics Hub", "2020 - 2022",
             ["Migrated Oracle DWH to Snowflake saving 40% in costs",
              "Orchestrated 200+ daily jobs using Apache Airflow",
              "Wrote dbt models for business reporting layer"]),
        ],
        "education": "B.Tech in Information Technology, BITS Pilani (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Aditya Patel",
        "email": "aditya.patel@gmail.com", "phone": "+91 9876509009",
        "summary": "Kusto (KQL) and Azure Data Explorer expert with 5+ years in cloud observability and log analytics.",
        "skills": "Kusto, KQL, Azure Data Explorer, Azure Monitor, Python, SQL, Grafana, Log Analytics, PowerBI, Azure",
        "experience": [
            ("Senior SRE / Analytics Engineer", "Microsoft India", "2021 - Present",
             ["Designed Kusto dashboards for real-time infra observability",
              "Wrote complex KQL queries across 50TB+ log datasets",
              "Built alerting systems reducing MTTR by 45%"]),
            ("Cloud Analytics Engineer", "Infosys", "2019 - 2021",
             ["Built Azure Monitor workbooks for cloud metrics",
              "Integrated Log Analytics with Grafana dashboards",
              "Automated daily KQL report generation with Python"]),
        ],
        "education": "B.Tech in Computer Science, NIT Warangal (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Riya Desai",
        "email": "riya.desai@gmail.com", "phone": "+91 9876510010",
        "summary": "Frontend developer with 3+ years in React, TypeScript, and modern CSS frameworks.",
        "skills": "React, TypeScript, JavaScript, HTML, CSS, Tailwind CSS, Redux, GraphQL, Jest, Git",
        "experience": [
            ("Frontend Engineer", "ProductLabs", "2022 - Present",
             ["Built complex multi-step form flows using React and Redux",
              "Improved Lighthouse performance score from 62 to 94",
              "Wrote unit tests with Jest achieving 90% code coverage"]),
            ("UI Developer", "Zeno Digital", "2021 - 2022",
             ["Created pixel-perfect React components from Figma designs",
              "Integrated GraphQL APIs with Apollo Client",
              "Reduced bundle size by 30% using code splitting"]),
        ],
        "education": "B.Tech in IT, Pune University (2017 - 2021)",
        "exp_years": 3,
    },
    {
        "name": "Siddharth Rao",
        "email": "siddharth.rao@gmail.com", "phone": "+91 9876511011",
        "summary": "Cybersecurity engineer with 6+ years in penetration testing, SIEM, and cloud security.",
        "skills": "Penetration Testing, SIEM, Splunk, Python, Azure Security, AWS IAM, Nessus, Burp Suite, Linux, Git",
        "experience": [
            ("Security Engineer", "SecureNet", "2020 - Present",
             ["Conducted 50+ penetration tests on web and mobile apps",
              "Built SIEM dashboards in Splunk for threat detection",
              "Implemented Azure Security Center policies across 200+ subscriptions"]),
            ("Junior Security Analyst", "CyberDefend", "2018 - 2020",
             ["Monitored SOC alerts and investigated incidents",
              "Automated threat intelligence feeds using Python scripts",
              "Performed vulnerability assessments using Nessus"]),
        ],
        "education": "B.Tech in Computer Science, Manipal University (2014 - 2018)",
        "exp_years": 6,
    },
    {
        "name": "Lakshmi Nair",
        "email": "lakshmi.nair@gmail.com", "phone": "+91 9876512012",
        "summary": "Product Manager with 5+ years turning complex technical requirements into shipped features.",
        "skills": "Product Management, Agile, Scrum, JIRA, Roadmapping, SQL, Data Analysis, Figma, Stakeholder Management",
        "experience": [
            ("Senior Product Manager", "SaaSify", "2021 - Present",
             ["Defined product roadmap for B2B SaaS platform with 50K users",
              "Collaborated with engineering and design to ship 3 major features quarterly",
              "Reduced churn by 18% via data-driven feature prioritization"]),
            ("Associate PM", "TechVista", "2019 - 2021",
             ["Managed backlog and sprint planning using JIRA",
              "Conducted user interviews and translated insights to requirements",
              "Analyzed product metrics using SQL and Mixpanel"]),
        ],
        "education": "MBA, IIM Ahmedabad (2017 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Arjun Bose",
        "email": "arjun.bose@gmail.com", "phone": "+91 9876513013",
        "summary": "Golang backend developer with 4+ years building high-performance APIs and distributed systems.",
        "skills": "Golang, gRPC, REST APIs, PostgreSQL, Redis, Kafka, Docker, Kubernetes, Prometheus, Git",
        "experience": [
            ("Backend Engineer (Go)", "Nexus Systems", "2022 - Present",
             ["Built gRPC services handling 50K requests per second",
              "Implemented distributed tracing using OpenTelemetry",
              "Reduced memory usage by 35% via goroutine pooling"]),
            ("Software Engineer", "SwiftCode", "2020 - 2022",
             ["Developed REST APIs in Go for logistics platform",
              "Designed PostgreSQL schemas for high-write workloads",
              "Implemented rate limiting middleware"]),
        ],
        "education": "B.Tech in Computer Science, IIIT Hyderabad (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Divya Menon",
        "email": "divya.menon@gmail.com", "phone": "+91 9876514014",
        "summary": "Salesforce developer with 5+ years in CRM customization, Apex development, and Lightning components.",
        "skills": "Salesforce, Apex, Lightning Web Components, SOQL, CRM, REST APIs, JavaScript, Git",
        "experience": [
            ("Senior Salesforce Developer", "CRMPros", "2021 - Present",
             ["Built custom Lightning components for Sales Cloud",
              "Integrated Salesforce with SAP via REST APIs",
              "Automated lead scoring using Apex triggers"]),
            ("Salesforce Developer", "TechForce", "2019 - 2021",
             ["Developed workflows and process builders",
              "Wrote complex SOQL queries for reporting",
              "Deployed changes using CI/CD with SFDX"]),
        ],
        "education": "B.E in IT, Osmania University (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Nikhil Gupta",
        "email": "nikhil.gupta@gmail.com", "phone": "+91 9876515015",
        "summary": "Android developer with 4+ years building high-quality apps in Kotlin and Java.",
        "skills": "Kotlin, Java, Android SDK, Jetpack Compose, MVVM, REST APIs, Firebase, Git, SQLite",
        "experience": [
            ("Senior Android Developer", "AppNation", "2022 - Present",
             ["Built fintech Android app with 500K+ downloads on Play Store",
              "Migrated XML layouts to Jetpack Compose reducing build time",
              "Integrated Firebase Analytics and Crashlytics"]),
            ("Android Developer", "MobileFirst", "2020 - 2022",
             ["Developed e-commerce Android app from scratch",
              "Implemented offline-first architecture with SQLite and sync",
              "Reduced app crash rate from 4% to 0.3%"]),
        ],
        "education": "B.Tech in IT, RGPV (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Pooja Thakur",
        "email": "pooja.thakur@gmail.com", "phone": "+91 9876516016",
        "summary": "Data Analyst with 3+ years extracting business insights using SQL, Python, and Power BI.",
        "skills": "SQL, Python, Power BI, Excel, Tableau, Pandas, Statistics, ETL, Azure Synapse",
        "experience": [
            ("Data Analyst", "InsightCo", "2022 - Present",
             ["Built Power BI dashboards tracking $10M in weekly revenue",
              "Wrote complex SQL queries joining 20+ tables for reports",
              "Automated weekly reporting using Python scripts"]),
            ("Junior Analyst", "AnalyticsEdge", "2021 - 2022",
             ["Cleaned and transformed raw data using Pandas",
              "Created Tableau visualizations for executive presentations",
              "Conducted statistical analysis for marketing campaigns"]),
        ],
        "education": "B.Sc in Statistics, Delhi University (2018 - 2021)",
        "exp_years": 3,
    },
    {
        "name": "Rahul Sinha",
        "email": "rahul.sinha@gmail.com", "phone": "+91 9876517017",
        "summary": "Site Reliability Engineer with 6+ years ensuring platform reliability, scalability, and observability.",
        "skills": "SRE, Python, Kubernetes, Prometheus, Grafana, PagerDuty, Terraform, AWS, Linux, Go",
        "experience": [
            ("Senior SRE", "CloudReliable", "2020 - Present",
             ["Maintained 99.99% uptime for platform with 10M+ users",
              "Implemented SLOs and error budgets reducing incidents by 40%",
              "Built automated remediation runbooks using Python and Go"]),
            ("DevOps Engineer", "TechStack", "2018 - 2020",
             ["Setup distributed tracing with Jaeger",
              "Wrote Terraform modules for AWS infrastructure",
              "Improved deployment frequency from weekly to daily"]),
        ],
        "education": "B.Tech in Computer Science, NIT Surathkal (2014 - 2018)",
        "exp_years": 6,
    },
    {
        "name": "Ishaan Kapoor",
        "email": "ishaan.kapoor@gmail.com", "phone": "+91 9876518018",
        "summary": "Blockchain developer with 4+ years in Solidity, smart contracts, and DeFi protocol development.",
        "skills": "Solidity, Ethereum, Web3.js, Smart Contracts, DeFi, Python, Node.js, Hardhat, IPFS, Git",
        "experience": [
            ("Blockchain Engineer", "DeFiCore", "2022 - Present",
             ["Developed ERC-20 and ERC-721 smart contracts with 0 critical bugs",
              "Built DeFi lending protocol managing $5M TVL",
              "Conducted smart contract audits using Slither and MythX"]),
            ("Web3 Developer", "ChainTech", "2020 - 2022",
             ["Integrated Web3.js with React dApp frontend",
              "Deployed contracts on Ethereum mainnet and testnets",
              "Built IPFS-based decentralized storage solutions"]),
        ],
        "education": "B.Tech in CS, IIIT Bangalore (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Tanvi Joshi",
        "email": "tanvi.joshi@gmail.com", "phone": "+91 9876519019",
        "summary": "QA Engineer with 5+ years in test automation, API testing, and performance testing.",
        "skills": "Selenium, Pytest, Postman, JMeter, Python, Java, REST APIs, CI/CD, TestNG, Git",
        "experience": [
            ("Senior QA Engineer", "QualityFirst", "2021 - Present",
             ["Built Selenium automation framework from scratch",
              "Achieved 80% test automation coverage across 3 products",
              "Conducted load testing with JMeter identifying critical bottlenecks"]),
            ("QA Engineer", "TestPro", "2019 - 2021",
             ["Wrote API test suites using Postman and RestAssured",
              "Integrated automated tests into Jenkins pipelines",
              "Maintained regression test suite of 2000+ test cases"]),
        ],
        "education": "B.Tech in IT, Symbiosis Institute (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Vivek Choudhary",
        "email": "vivek.c@gmail.com", "phone": "+91 9876520020",
        "summary": "Embedded systems engineer with 6+ years in firmware development, RTOS, and IoT solutions.",
        "skills": "C, C++, RTOS, FreeRTOS, STM32, IoT, MQTT, Bluetooth, WiFi, Embedded Linux, Git",
        "experience": [
            ("Senior Embedded Engineer", "IoTech", "2020 - Present",
             ["Developed firmware for industrial IoT gateway devices",
              "Reduced boot time from 12s to 3s through boot optimization",
              "Implemented secure OTA update mechanism using MQTT"]),
            ("Embedded Developer", "Microtech", "2018 - 2020",
             ["Wrote bare-metal drivers for STM32 microcontrollers",
              "Ported FreeRTOS to custom hardware",
              "Debugged hardware issues using oscilloscope and JTAG"]),
        ],
        "education": "B.Tech in Electronics, NIT Calicut (2014 - 2018)",
        "exp_years": 6,
    },
    {
        "name": "Neha Agarwal",
        "email": "neha.agarwal@gmail.com", "phone": "+91 9876521021",
        "summary": "Business Intelligence developer with 5+ years in SQL, Power BI, and enterprise data warehousing.",
        "skills": "SQL, Power BI, SSAS, SSIS, SSRS, Azure Synapse, Python, DAX, MDX, Snowflake",
        "experience": [
            ("BI Developer", "DataInsight", "2021 - Present",
             ["Built enterprise Power BI reports for 500+ business users",
              "Designed Azure Synapse data warehouse replacing legacy DWH",
              "Optimized DAX queries reducing report load time by 50%"]),
            ("SQL Developer", "ReportMasters", "2019 - 2021",
             ["Wrote complex stored procedures for financial reporting",
              "Developed SSIS packages for daily ETL workflows",
              "Created SSRS reports for executive dashboards"]),
        ],
        "education": "B.Com with CS minor, Mumbai University (2015 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Harsh Vardhan",
        "email": "harsh.v@gmail.com", "phone": "+91 9876522022",
        "summary": "iOS developer with 4+ years shipping Swift apps with a focus on performance and clean architecture.",
        "skills": "Swift, SwiftUI, UIKit, Combine, CoreData, REST APIs, Firebase, Xcode, Git, TestFlight",
        "experience": [
            ("iOS Developer", "AppStudio", "2022 - Present",
             ["Rebuilt legacy UIKit app in SwiftUI improving maintainability",
              "Published 3 apps with combined 200K+ downloads",
              "Implemented offline sync using CoreData and CloudKit"]),
            ("Junior iOS Developer", "Appify", "2020 - 2022",
             ["Developed features for healthcare iOS application",
              "Integrated HealthKit APIs for fitness tracking",
              "Wrote unit tests with XCTest maintaining 75% coverage"]),
        ],
        "education": "B.Tech in CS, Jadavpur University (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Kavya Reddy",
        "email": "kavya.reddy@gmail.com", "phone": "+91 9876523023",
        "summary": "AI/ML researcher with 3+ years in computer vision, image classification, and object detection.",
        "skills": "Python, PyTorch, OpenCV, YOLO, CNN, Transfer Learning, Docker, CUDA, NumPy, Pandas",
        "experience": [
            ("Computer Vision Engineer", "VisionAI", "2023 - Present",
             ["Built YOLO-based real-time object detection for retail analytics",
              "Achieved 95% mAP on custom dataset using transfer learning",
              "Optimized inference pipeline for edge deployment"]),
            ("ML Research Intern", "AI Labs India", "2022 - 2023",
             ["Reproduced SOTA image segmentation paper results",
              "Trained CNNs on medical imaging datasets",
              "Published paper on augmentation strategies"]),
        ],
        "education": "M.Tech in AI, IIT Kharagpur (2020 - 2022)",
        "exp_years": 3,
    },
    {
        "name": "Sameer Kulkarni",
        "email": "sameer.k@gmail.com", "phone": "+91 9876524024",
        "summary": "SAP developer with 7+ years in ABAP, SAP HANA, and S/4HANA implementations.",
        "skills": "SAP ABAP, SAP HANA, S/4HANA, Fiori, BAPI, BADI, SAP BTP, SQL, SAP MM/SD, Git",
        "experience": [
            ("SAP Senior Consultant", "SAPExperts", "2020 - Present",
             ["Led S/4HANA greenfield implementation for manufacturing client",
              "Developed custom ABAP programs reducing manual effort by 60%",
              "Built SAP Fiori apps for mobile inventory management"]),
            ("SAP Developer", "TechConsult", "2017 - 2020",
             ["Wrote ABAP reports and enhancements for SAP MM module",
              "Optimized HANA queries for analytics use cases",
              "Conducted unit and integration testing for SAP projects"]),
        ],
        "education": "B.E in CS, Pune University (2013 - 2017)",
        "exp_years": 7,
    },
    {
        "name": "Shreya Banerjee",
        "email": "shreya.b@gmail.com", "phone": "+91 9876525025",
        "summary": "ML Ops engineer with 3+ years automating model training, deployment, and monitoring pipelines.",
        "skills": "MLflow, Kubeflow, Docker, Kubernetes, Python, Terraform, AWS SageMaker, Prometheus, CI/CD, Git",
        "experience": [
            ("MLOps Engineer", "ModelScale", "2023 - Present",
             ["Built end-to-end ML platform using Kubeflow on GKE",
              "Automated model retraining triggers using Airflow DAGs",
              "Implemented model drift detection reducing incidents by 30%"]),
            ("Data Scientist", "PredictAI", "2021 - 2023",
             ["Trained and evaluated classification models using Scikit-learn",
              "Deployed models to AWS SageMaker endpoints",
              "Monitored model performance using custom Prometheus metrics"]),
        ],
        "education": "M.Tech in CS, BITS Hyderabad (2019 - 2021)",
        "exp_years": 3,
    },
    {
        "name": "Akash Tiwari",
        "email": "akash.tiwari@gmail.com", "phone": "+91 9876526026",
        "summary": "Network engineer with 6+ years in enterprise networking, SDN, and cloud networking.",
        "skills": "Cisco, BGP, OSPF, SDN, Python, AWS VPC, Azure Networking, Wireshark, Linux, Terraform",
        "experience": [
            ("Senior Network Engineer", "NetSolutions", "2020 - Present",
             ["Designed multi-region AWS VPC architecture for 99.99% availability",
              "Migrated MPLS network to SD-WAN reducing costs by 35%",
              "Automated network provisioning using Python and Ansible"]),
            ("Network Engineer", "ConnectTech", "2018 - 2020",
             ["Configured BGP and OSPF routing for enterprise WAN",
              "Troubleshot network performance using Wireshark",
              "Implemented network security policies and VLANs"]),
        ],
        "education": "B.Tech in ECE, Amrita University (2014 - 2018)",
        "exp_years": 6,
    },
    {
        "name": "Prathyusha Rao",
        "email": "prathyusha.r@gmail.com", "phone": "+91 9876527027",
        "summary": "Scrum Master and Agile coach with 5+ years driving delivery excellence in cross-functional teams.",
        "skills": "Scrum, Kanban, JIRA, Confluence, Agile, SAFe, Stakeholder Management, Risk Management, Retrospectives",
        "experience": [
            ("Senior Scrum Master", "AgilePro", "2021 - Present",
             ["Coached 4 scrum teams across 2 products in SAFe framework",
              "Improved team velocity by 30% through process improvements",
              "Facilitated PI Planning for 80+ participants"]),
            ("Scrum Master", "SprintForce", "2019 - 2021",
             ["Removed impediments enabling consistent sprint delivery",
              "Ran retrospectives identifying and fixing recurring issues",
              "Maintained JIRA boards and sprint reports"]),
        ],
        "education": "MBA, XLRI Jamshedpur (2017 - 2019)",
        "exp_years": 5,
    },
    {
        "name": "Ankit Sharma",
        "email": "ankit.sharma@gmail.com", "phone": "+91 9876528028",
        "summary": "Database administrator with 8+ years managing Oracle, PostgreSQL, and Azure SQL databases.",
        "skills": "Oracle, PostgreSQL, MySQL, Azure SQL, SQL Server, PL/SQL, Performance Tuning, Backup, Python, Linux",
        "experience": [
            ("Senior DBA", "DataGuard", "2019 - Present",
             ["Managed 50+ production databases with 99.99% SLA",
              "Tuned slow queries reducing average response time by 65%",
              "Implemented automated backup and DR strategies"]),
            ("DBA", "DataTech", "2016 - 2019",
             ["Performed schema design and normalization for ERP system",
              "Wrote PL/SQL procedures for batch processing",
              "Migrated on-prem Oracle DB to Azure SQL"]),
        ],
        "education": "B.Tech in IT, JNTU Hyderabad (2012 - 2016)",
        "exp_years": 8,
    },
    {
        "name": "Megha Singh",
        "email": "megha.singh@gmail.com", "phone": "+91 9876529029",
        "summary": "UX Designer with 4+ years creating user-centered digital products using Figma and design systems.",
        "skills": "Figma, Adobe XD, UX Research, Wireframing, Prototyping, Design Systems, HTML, CSS, Usability Testing",
        "experience": [
            ("Senior UX Designer", "DesignFirst", "2022 - Present",
             ["Led redesign of B2C app increasing user retention by 22%",
              "Built and maintained company-wide design system in Figma",
              "Conducted 30+ usability testing sessions"]),
            ("UX Designer", "CreativeStudio", "2020 - 2022",
             ["Designed wireframes and prototypes for fintech app",
              "Collaborated with developers to ensure pixel-perfect implementation",
              "Created user journey maps and personas"]),
        ],
        "education": "B.Des, NID Ahmedabad (2016 - 2020)",
        "exp_years": 4,
    },
    {
        "name": "Rajesh Kumar",
        "email": "rajesh.kumar@gmail.com", "phone": "+91 9876530030",
        "summary": "Technical writer with 5+ years producing API docs, developer guides, and knowledge base articles.",
        "skills": "Technical Writing, API Documentation, Swagger, Markdown, DITA, Confluence, Git, Python basics, Postman",
        "experience": [
            ("Senior Technical Writer", "DocCraft", "2021 - Present",
             ["Wrote and maintained REST API documentation for 150+ endpoints",
              "Reduced support tickets by 40% through improved developer docs",
              "Built docs-as-code pipeline using MkDocs and GitHub Actions"]),
            ("Technical Writer", "WriteCode", "2019 - 2021",
             ["Created onboarding guides for developer platform",
              "Maintained Confluence knowledge base of 500+ articles",
              "Collaborated with engineers to document internal APIs"]),
        ],
        "education": "B.A in English Literature, Delhi University (2015 - 2019)",
        "exp_years": 5,
    },
]


# ─── DOCX generator ───────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 60)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def create_resume_docx(data, output_path):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(54)

    # Header
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(data["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    contact_p = doc.add_paragraph()
    contact_p.add_run(f"Email: {data['email']}  |  Phone: {data['phone']}")
    contact_p.runs[0].font.size = Pt(10)

    add_divider(doc)

    # Professional Summary
    add_heading(doc, "PROFESSIONAL SUMMARY", 2)
    doc.add_paragraph(data["summary"])

    add_divider(doc)

    # Skills
    add_heading(doc, "SKILLS", 2)
    doc.add_paragraph(data["skills"])

    add_divider(doc)

    # Experience
    add_heading(doc, "EXPERIENCE", 2)
    for title, company, period, bullets in data["experience"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{title} – {company} ({period})")
        r.bold = True
        r.font.size = Pt(11)
        for bullet in bullets:
            doc.add_paragraph(f"• {bullet}", style="List Bullet")

    add_divider(doc)

    # Education
    add_heading(doc, "EDUCATION", 2)
    doc.add_paragraph(data["education"])

    doc.save(output_path)


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"Generating {len(RESUMES)} resumes in '{OUTPUT_DIR}/'...\n")

    for r in RESUMES:
        safe_name = r["name"].lower().replace(" ", "_")
        out_path = os.path.join(OUTPUT_DIR, f"{safe_name}.docx")
        create_resume_docx(r, out_path)
        print(f"  ✅ {r['name']} → {safe_name}.docx")

    print(f"\n🎉 Done! {len(RESUMES)} resumes saved to '{OUTPUT_DIR}/'")
    print("   Now run: py resume_rag.py")